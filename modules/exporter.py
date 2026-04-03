"""Export generated MCQs as .docx or Moodle-compatible XML."""

import os
import uuid
import xml.etree.ElementTree as ET
from xml.dom import minidom

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

EXPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads")
OPTION_LABELS = ["A", "B", "C", "D"]
BLOOM_COLORS = {
    "Remember": RGBColor(0x27, 0xAE, 0x60),   # green
    "Apply":    RGBColor(0x29, 0x80, 0xB9),   # blue
    "Analyse":  RGBColor(0x8E, 0x44, 0xAD),   # purple
}


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def export_docx(questions: list[dict], result_id: str) -> str:
    doc = Document()

    # Title
    title = doc.add_heading("EduQuiz Forge — Generated Question Paper", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    current_level = None
    for q in questions:
        bloom = q.get("bloom_level", "")

        # Section header when Bloom level changes
        if bloom != current_level:
            current_level = bloom
            hdr = doc.add_heading(f"Section: {bloom} Level", level=2)
            color = BLOOM_COLORS.get(bloom, RGBColor(0, 0, 0))
            for run in hdr.runs:
                run.font.color.rgb = color
            doc.add_paragraph()

        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{q.get('number', '')}. {q['question']}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        # Options
        for i, opt in enumerate(q["options"]):
            marker = "→" if i == q["answer_index"] else " "
            doc.add_paragraph(f"   {OPTION_LABELS[i]}) {opt}", style="List Bullet")

        # Answer & explanation
        correct_label = OPTION_LABELS[q["answer_index"]]
        ans_para = doc.add_paragraph()
        ans_run = ans_para.add_run(f"Answer: {correct_label}   |   {q['explanation']}")
        ans_run.italic = True
        ans_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        ans_run.font.size = Pt(9)

        doc.add_paragraph()

    out_path = os.path.join(EXPORTS_DIR, f"eduquiz_{result_id}.docx")
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Moodle XML export
# ---------------------------------------------------------------------------

def _indent_xml(elem: ET.Element, level: int = 0) -> None:
    """In-place pretty-print indentation for ElementTree."""
    indent = "\n" + "  " * level
    if len(elem):
        if not elem.text or not elem.text.strip():
            elem.text = indent + "  "
        if not elem.tail or not elem.tail.strip():
            elem.tail = indent
        for child in elem:
            _indent_xml(child, level + 1)
        if not child.tail or not child.tail.strip():
            child.tail = indent
    else:
        if level and (not elem.tail or not elem.tail.strip()):
            elem.tail = indent


def export_moodle_xml(questions: list[dict], result_id: str) -> str:
    quiz = ET.Element("quiz")

    for q in questions:
        question_el = ET.SubElement(quiz, "question", type="multichoice")

        name_el = ET.SubElement(question_el, "name")
        text_el = ET.SubElement(name_el, "text")
        text_el.text = f"Q{q.get('number', '')} [{q.get('bloom_level','')}]"

        qtext_el = ET.SubElement(question_el, "questiontext", format="html")
        qt_text = ET.SubElement(qtext_el, "text")
        qt_text.text = f"<![CDATA[{q['question']}]]>"

        general_fb = ET.SubElement(question_el, "generalfeedback", format="html")
        gf_text = ET.SubElement(general_fb, "text")
        gf_text.text = f"<![CDATA[{q['explanation']}]]>"

        ET.SubElement(question_el, "defaultgrade").text = "1"
        ET.SubElement(question_el, "penalty").text = "0"
        ET.SubElement(question_el, "hidden").text = "0"
        ET.SubElement(question_el, "single").text = "true"
        ET.SubElement(question_el, "shuffleanswers").text = "true"

        for i, opt in enumerate(q["options"]):
            fraction = "100" if i == q["answer_index"] else "0"
            ans_el = ET.SubElement(question_el, "answer", fraction=fraction, format="html")
            ans_text = ET.SubElement(ans_el, "text")
            ans_text.text = f"<![CDATA[{opt}]]>"
            fb_el = ET.SubElement(ans_el, "feedback", format="html")
            fb_text = ET.SubElement(fb_el, "text")
            fb_text.text = "" if i != q["answer_index"] else f"<![CDATA[{q['explanation']}]]>"

        # Bloom's taxonomy tag
        tags_el = ET.SubElement(question_el, "tags")
        tag_el = ET.SubElement(tags_el, "tag")
        tag_text = ET.SubElement(tag_el, "text")
        tag_text.text = q.get("bloom_level", "")

    _indent_xml(quiz)
    tree = ET.ElementTree(quiz)

    out_path = os.path.join(EXPORTS_DIR, f"eduquiz_{result_id}.xml")
    tree.write(out_path, encoding="utf-8", xml_declaration=True)
    return out_path
